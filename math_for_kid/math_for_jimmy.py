import sys
import numpy as np
import time
import docx
from docx.shared import RGBColor
from docx.shared import Pt

#change these limit to change difficulty
ADD_LIMIT = 25
MULTIPLY_LIMIT = 9
ALLOW_NEGATIVE = False #this only affects add and sub

#word related parameters
colors = [RGBColor(0,0,0),RGBColor(255,0,0),RGBColor(0,128,0),RGBColor(0,0,255),RGBColor(255,128,0),RGBColor(255,0,255)]
max_color = len(colors)
font_size = 30
number_of_questions = 13

np.random.seed(int(time.time()))

def generate_question(add,mult):
    """generate one question string with add and mult limit"""
    op_array = ["+","-","x"]
    operator = np.random.randint(3)
    if operator == 2:
        op1 = np.random.randint(mult)+1
        op2 = np.random.randint(mult)+1
    else:
        #negative is allowed, add and subtract are similar
        if ALLOW_NEGATIVE:
            op1 = np.random.randint(add)+1
            op2 = np.random.randint(add)+1
            op1 = op1 if np.random.randint(2) else -op1
            op2 = op2 if np.random.randint(2) else -op2
        #negative not allowed, subtract need to consider operands values
        else:
            op1 = np.random.randint(add)+1
            #addition case, operands can be anything within range
            if operator == 0:
                op2 = np.random.randint(add)+1
            #subtract case, second operand should be smaller than first one
            else:
                op2 = np.random.randint(op1)+1
    return str(op1)+op_array[operator]+str(op2)+" ="

def generate_question_double(add,mult):
    """Generate one line with two questions, looking like a two column page"""
    q1 = generate_question(add,mult)
    q2 = generate_question(add,mult)
    return q1+(13-len(q1))*' '+q2

def create_test():
    doc = docx.Document()
    for i in range(number_of_questions):
        p = doc.add_paragraph()
        question = generate_question_double(ADD_LIMIT,MULTIPLY_LIMIT)
        for j in range(len(question)):
            run = p.add_run(question[j])
            run.bold = True
            run.font.name = "Courier New"
            run.font.size = Pt(font_size)
            run.font.color.rgb = colors[np.random.randint(max_color)]
    t = time.localtime()
    year = str(t.tm_year)
    month = str(t.tm_mon)
    day = str(t.tm_mday)
    filename = year+"_"+month+"_"+day+".docx"
    doc.save(filename)
        
if __name__ == "__main__":
    create_test()
    
